"""
Parsers for pre-authored quiz file formats.

Format 2 — Hash TXT:
    1<TAB>Question text
    A<TAB>#Correct answer  (# marks the correct option)
    B<TAB>Wrong answer
    C<TAB>Wrong answer
    D<TAB>Wrong answer

Format 3 — DOCX Table:
    Each table = one question
    First row (first cell) = question text
    Remaining rows = answer options; FIRST option row is always correct

Both parsers return a list of normalised question dicts that match the
structure expected by QuizSession:
    {
        "question": str,
        "options":  {"A": str, "B": str, "C": str, "D": str},
        "correct":  "A" | "B" | "C" | "D",
        "explanation": "",
    }

Answer options are shuffled so the correct answer is never predictably
in position A.
"""

import io
import random
from docx import Document


# ---------------------------------------------------------------------------
# Shared helpers
# ---------------------------------------------------------------------------

def _shuffle_options(options: dict, correct_text: str) -> tuple[dict, str]:
    """Shuffle option values, reassign keys A-D, return (new_options, new_correct_key)."""
    keys = ["A", "B", "C", "D"]
    values = list(options.values())
    random.shuffle(values)
    new_options = dict(zip(keys, values))
    new_correct = next(k for k, v in new_options.items() if v == correct_text)
    return new_options, new_correct


def _build_question(question_text: str, options_raw: dict, correct_text: str) -> dict:
    """Build a normalised question dict with shuffled options."""
    options, correct_key = _shuffle_options(options_raw, correct_text)
    return {
        "question": question_text.strip(),
        "options": options,
        "correct": correct_key,
        "explanation": "",
    }


# ---------------------------------------------------------------------------
# Format 2 — Hash TXT
# ---------------------------------------------------------------------------

def _is_hash_format(text: str) -> bool:
    """Return True if the text looks like Format 2 (tab-separated with # markers)."""
    lines = [l for l in text.splitlines() if l.strip()]
    hash_count = sum(1 for l in lines if "\t" in l and l.split("\t", 1)[-1].startswith("#"))
    return hash_count >= 2


def parse_hash_txt(text: str) -> list[dict]:
    """
    Parse Format 2 (hash TXT).  Blocks are separated by blank lines.
    Each block:
        <number>TAB<question>
        <letter>TAB#<correct answer>
        <letter>TAB<wrong answer>
        ...
    """
    questions = []
    # Split into blocks on blank lines
    raw_blocks = []
    current: list[str] = []
    for line in text.splitlines():
        stripped = line.strip()
        if stripped:
            current.append(stripped)
        else:
            if current:
                raw_blocks.append(current)
                current = []
    if current:
        raw_blocks.append(current)

    for block in raw_blocks:
        # First line should be <number/label><TAB><question>
        if "\t" not in block[0]:
            continue
        question_text = block[0].split("\t", 1)[1].strip()

        options_raw: dict[str, str] = {}
        correct_text: str | None = None

        for line in block[1:]:
            if "\t" not in line:
                continue
            parts = line.split("\t", 1)
            letter = parts[0].strip().upper()
            value = parts[1].strip()
            if len(letter) != 1 or letter not in "ABCDEFGH":
                continue
            if value.startswith("#"):
                answer_text = value[1:].strip()
                correct_text = answer_text
                options_raw[letter] = answer_text
            else:
                options_raw[letter] = value

        # Need at least 2 options and a known correct answer
        if not question_text or correct_text is None or len(options_raw) < 2:
            continue

        # Pad to exactly 4 options if fewer were provided (edge case)
        all_keys = ["A", "B", "C", "D"]
        while len(options_raw) < 4:
            for k in all_keys:
                if k not in options_raw:
                    options_raw[k] = "—"
                    break

        # Keep only first 4 keys
        trimmed = {k: options_raw[k] for k in all_keys if k in options_raw}
        if len(trimmed) < 4:
            continue

        questions.append(_build_question(question_text, trimmed, correct_text))

    return questions


# ---------------------------------------------------------------------------
# Format 3 — DOCX Table
# ---------------------------------------------------------------------------

def _is_table_format(file_bytes: bytes) -> bool:
    """Return True if the DOCX has at least one table with ≥2 rows."""
    try:
        doc = Document(io.BytesIO(file_bytes))
        return any(len(t.rows) >= 2 for t in doc.tables)
    except Exception:
        return False


def parse_docx_tables(file_bytes: bytes) -> list[dict]:
    """
    Parse Format 3 (DOCX table).
    Each table = one question.
    Row 0 col 0 = question text.
    Remaining rows = answer options; first option row = correct answer.
    """
    doc = Document(io.BytesIO(file_bytes))
    questions = []

    for table in doc.tables:
        rows = table.rows
        if len(rows) < 2:
            continue

        # Question text from first row
        question_text = " ".join(
            cell.text.strip() for cell in rows[0].cells if cell.text.strip()
        ).strip()
        if not question_text:
            continue

        # Collect answer option texts (first option = correct)
        option_texts = []
        for row in rows[1:]:
            text = " ".join(cell.text.strip() for cell in row.cells if cell.text.strip()).strip()
            if text:
                option_texts.append(text)

        if len(option_texts) < 2:
            continue

        correct_text = option_texts[0]

        # Ensure exactly 4 options (pad or truncate)
        while len(option_texts) < 4:
            option_texts.append("—")
        option_texts = option_texts[:4]

        options_raw = {"A": option_texts[0], "B": option_texts[1],
                       "C": option_texts[2], "D": option_texts[3]}

        questions.append(_build_question(question_text, options_raw, correct_text))

    return questions


# ---------------------------------------------------------------------------
# Auto-detection entry point
# ---------------------------------------------------------------------------

def detect_and_parse(
    file_bytes: bytes,
    file_name: str,
    raw_text: str,
) -> tuple[str, list[dict] | str]:
    """
    Detect the quiz format and parse if possible.

    Returns:
        ("parsed",  list[dict])  — pre-parsed questions, skip AI
        ("ai",      str)         — raw text to feed to Gemini
    """
    name_lower = file_name.lower()

    # TXT file — check for hash format
    if name_lower.endswith(".txt"):
        if _is_hash_format(raw_text):
            questions = parse_hash_txt(raw_text)
            if questions:
                return "parsed", questions
        # TXT without hash format — feed as plain text to AI
        return "ai", raw_text

    # DOCX — check for table format first
    if name_lower.endswith(".docx"):
        if _is_table_format(file_bytes):
            questions = parse_docx_tables(file_bytes)
            if questions:
                return "parsed", questions
        # DOCX without tables — fall through to AI
        return "ai", raw_text

    # PDF, PPTX, everything else → AI
    return "ai", raw_text
